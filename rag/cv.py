"""CV extraction and role recommendations using the Career Tree RAG."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path
from typing import Any, Dict, List

from rag.config import CV_CACHE_DIR

MAX_CV_TEXT_CHARS = 15_000
DEFAULT_CV_EXTRACTION_MODEL = "qwen2.5-0.5b"


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    texts: List[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        texts.append(t)
    return "\n".join(texts)


def _read_docx(path: Path) -> str:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    # basic tables support (cells often hold bullet-like text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs if p.text)
                if cell_text:
                    parts.append(cell_text)
    return "\n".join(parts)


def read_cv_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = _read_text_file(path)
    elif suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix in {".docx"}:
        text = _read_docx(path)
    else:
        text = _read_text_file(path)

    if not text.strip():
        raise SystemExit(
            f"No readable text found in CV: {path}. "
            "For scanned PDFs, convert to text or use OCR first."
        )
    return text


CV_JSON_SCHEMA = {
    "name": "string",
    "contact": {"email": "string", "phone": "string", "location": "string"},
    "summary": "string",
    "years_experience": "number",
    "skills": ["string"],
    "education": [
        {
            "degree": "string",
            "field": "string",
            "institution": "string",
            "start_year": "number|null",
            "end_year": "number|null",
        }
    ],
    "experience": [
        {
            "company": "string",
            "role": "string",
            "start": "string",
            "end": "string",
            "highlights": ["string"],
            "skills": ["string"],
        }
    ],
    "projects": [
        {
            "name": "string",
            "description": "string",
            "skills": ["string"],
        }
    ],
    "certifications": ["string"],
    "languages": ["string"],
}


def _json_schema_str() -> str:
    return json.dumps(CV_JSON_SCHEMA, ensure_ascii=False, indent=2)


def _normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _truncate_cv_text(text: str, max_chars: int = MAX_CV_TEXT_CHARS) -> str:
    cleaned = text.strip()
    if len(cleaned) <= max_chars:
        return cleaned
    return cleaned[: max_chars - 3].rstrip() + "..."


def _parse_cv_json(output: str) -> Dict[str, Any]:
    raw = output or "{}"
    first = raw.find("{")
    last = raw.rfind("}")
    if first != -1 and last != -1 and last > first:
        raw = raw[first : last + 1]
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise ValueError(f"CV extraction returned invalid JSON: {exc}") from exc
    if not isinstance(data, dict):
        raise ValueError("CV extraction did not return a JSON object")
    return data


def _normalize_cv_data(data: Dict[str, Any]) -> Dict[str, Any]:
    if "skills" in data and isinstance(data["skills"], list):
        data["skills"] = sorted(
            {_normalize_ws(skill) for skill in data["skills"] if isinstance(skill, str) and skill.strip()}
        )
    return data


def _extract_with_chat(chat, text: str, *, strict: bool = False) -> Dict[str, Any]:
    system = (
        "You are a precise CV information extractor. Return ONLY valid JSON matching the provided schema. "
        "Do not include markdown fences or commentary."
    )
    if strict:
        system += " Output must be a single JSON object and nothing else."

    user = (
        "Extract the candidate information from the CV text below.\n\n"
        f"Return JSON strictly matching this schema (omit unknown fields, use nulls sparingly):\n\n"
        f"{_json_schema_str()}\n\n"
        f"CV Text:\n{text}"
    )
    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]
    completion = chat.complete_chat(messages)
    output = completion.choices[0].message.content or "{}"
    return _parse_cv_json(output)


def extract_cv_structured(
    text: str,
    model_alias: str | None = None,
    *,
    chat=None,
) -> Dict[str, Any]:
    """Use a local LLM to extract structured CV data as JSON.

    Pass an already-loaded `chat` client to avoid loading the model twice.
    """
    cv_text = _truncate_cv_text(text)
    own_model = None

    if chat is None:
        from rag.llm import load_model, make_chat

        model_alias = model_alias or DEFAULT_CV_EXTRACTION_MODEL
        own_model = load_model(model_alias)
        chat = make_chat(own_model, temperature=0.1, max_tokens=2000)

    try:
        try:
            data = _extract_with_chat(chat, cv_text)
        except ValueError:
            data = _extract_with_chat(chat, cv_text, strict=True)
    finally:
        if own_model is not None:
            own_model.unload()

    return _normalize_cv_data(data)


def _unique_strings(items: Any, limit: int | None = None) -> List[str]:
    seen: set[str] = set()
    out: List[str] = []
    if not isinstance(items, list):
        return out
    for item in items:
        if not isinstance(item, str) or not item.strip():
            continue
        normalized = _normalize_ws(item)
        key = normalized.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(normalized)
        if limit is not None and len(out) >= limit:
            break
    return out


def _cv_cache_path(cv_path: Path, model_alias: str) -> Path:
    digest = hashlib.sha256(cv_path.read_bytes()).hexdigest()[:16]
    safe_model = re.sub(r"[^a-z0-9_-]+", "-", model_alias.lower()).strip("-") or "default"
    return CV_CACHE_DIR / f"{cv_path.stem}-{digest}-{safe_model}.json"


def load_or_extract_cv(
    cv_path: Path,
    *,
    model_alias: str | None = None,
    chat=None,
    use_cache: bool = True,
) -> Dict[str, Any]:
    """Read a CV file and return structured JSON, using cache when available."""
    text = read_cv_text(cv_path)
    model_alias = model_alias or DEFAULT_CV_EXTRACTION_MODEL
    cache_file = _cv_cache_path(cv_path, model_alias)

    if use_cache and cache_file.exists():
        try:
            cached = json.loads(cache_file.read_text(encoding="utf-8"))
            if isinstance(cached, dict) and cached:
                print(f"Using cached CV profile: {cache_file.name}")
                return _normalize_cv_data(cached)
        except json.JSONDecodeError:
            pass

    data = extract_cv_structured(text, model_alias=model_alias, chat=chat)
    if use_cache:
        CV_CACHE_DIR.mkdir(parents=True, exist_ok=True)
        cache_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"Cached CV profile: {cache_file.name}")
    return data


def _build_query_from_cv(cv: Dict[str, Any]) -> str:
    parts: List[str] = []

    summary = cv.get("summary")
    if summary:
        parts.append(_normalize_ws(str(summary)))

    skills = _unique_strings(cv.get("skills") or [], 20)
    if skills:
        parts.append("skills: " + ", ".join(skills))

    edu = cv.get("education") or []
    if isinstance(edu, list):
        edu_bits: List[str] = []
        for item in edu:
            if not isinstance(item, dict):
                continue
            field = item.get("field")
            degree = item.get("degree")
            if field:
                edu_bits.append(str(field))
            elif degree:
                edu_bits.append(str(degree))
        edu_bits = _unique_strings(edu_bits, 3)
        if edu_bits:
            parts.append("education: " + ", ".join(edu_bits))

    exp = cv.get("experience") or []
    if isinstance(exp, list):
        roles: List[str] = []
        highlights: List[str] = []
        exp_skills: List[str] = []
        for item in exp:
            if not isinstance(item, dict):
                continue
            if item.get("role"):
                roles.append(str(item["role"]))
            highlights.extend(item.get("highlights") or [])
            exp_skills.extend(item.get("skills") or [])
        roles = _unique_strings(roles, 4)
        if roles:
            parts.append("roles: " + ", ".join(roles))
        highlight_bits = _unique_strings(highlights, 4)
        if highlight_bits:
            parts.append("experience highlights: " + "; ".join(highlight_bits))
        exp_skill_bits = _unique_strings(exp_skills, 8)
        if exp_skill_bits:
            parts.append("role skills: " + ", ".join(exp_skill_bits))

    projects = cv.get("projects") or []
    if isinstance(projects, list):
        project_bits: List[str] = []
        project_skills: List[str] = []
        for item in projects:
            if not isinstance(item, dict):
                continue
            name = item.get("name")
            description = item.get("description")
            if name:
                project_bits.append(str(name))
            elif description:
                project_bits.append(_normalize_ws(str(description))[:120])
            project_skills.extend(item.get("skills") or [])
        project_bits = _unique_strings(project_bits, 3)
        if project_bits:
            parts.append("projects: " + ", ".join(project_bits))
        project_skill_bits = _unique_strings(project_skills, 6)
        if project_skill_bits:
            parts.append("project skills: " + ", ".join(project_skill_bits))

    certs = _unique_strings(cv.get("certifications") or [], 5)
    if certs:
        parts.append("certifications: " + ", ".join(certs))

    return "; ".join(parts)


# Public alias for use by other modules
def build_query_from_cv(cv: Dict[str, Any]) -> str:  # noqa: D401
    """Build a short retrieval query string from CV JSON (skills/roles/summary)."""
    return _build_query_from_cv(cv)


def compact_cv_summary(cv: Dict[str, Any], max_skills: int = 12) -> str:
    """Compact, human-readable CV summary for prompts (keeps it short)."""
    parts: List[str] = []
    years = cv.get("years_experience")
    if years:
        parts.append(f"experience: {years}y")

    skills = _unique_strings(cv.get("skills") or [], max_skills)
    if skills:
        parts.append("skills: " + ", ".join(skills))

    exp = cv.get("experience") or []
    if isinstance(exp, list):
        roles = [
            str(e.get("role"))
            for e in exp
            if isinstance(e, dict) and e.get("role")
        ]
        roles = _unique_strings(roles, 3)
        if roles:
            parts.append("roles: " + ", ".join(roles))
        highlights = []
        for item in exp:
            if isinstance(item, dict):
                highlights.extend(item.get("highlights") or [])
        highlight_bits = _unique_strings(highlights, 2)
        if highlight_bits:
            parts.append("highlights: " + "; ".join(highlight_bits))

    edu = cv.get("education") or []
    if isinstance(edu, list):
        fields = [
            str(e.get("field") or e.get("degree"))
            for e in edu
            if isinstance(e, dict) and (e.get("field") or e.get("degree"))
        ]
        fields = _unique_strings(fields, 2)
        if fields:
            parts.append("education: " + ", ".join(fields))

    certs = _unique_strings(cv.get("certifications") or [], 3)
    if certs:
        parts.append("certifications: " + ", ".join(certs))

    return "; ".join(parts)


def recommend_from_cv(cv: Dict[str, Any], top_k: int = 5, layer: str | None = None) -> Dict[str, Any]:
    from rag.retrieve import retrieve

    query = _build_query_from_cv(cv)
    if not query:
        # fallback: ask for AI/ML roles based on skills
        skills = cv.get("skills") or []
        query = "AI/ML roles for: " + ", ".join(skills) if skills else "AI/ML roles"

    hits = retrieve(query, top_k=top_k, layer=layer)
    return {"query": query, "hits": hits}


def _plan_recommendations(
    cv: Dict[str, Any],
    hits: List[Dict[str, Any]],
    model_alias: str | None = None,
    *,
    chat=None,
) -> str:
    from rag.ask import _llm_context_from_hits

    own_model = None
    if chat is None:
        from rag.llm import DEFAULT_MODEL, load_model, make_chat

        model_alias = model_alias or DEFAULT_MODEL
        own_model = load_model(model_alias)
        chat = make_chat(own_model, temperature=0.2, max_tokens=1200)

    try:
        context = _llm_context_from_hits(hits)
        cv_preview = json.dumps(cv, ensure_ascii=False, indent=2)
        system = (
            "You are a career advisor. Using the user's CV data and the provided role/context snippets,"
            " recommend the top 5 fitting roles from the context. For each role provide:"
            " 1) role name, 2) fit score 0-100, 3) why it matches, 4) skill gaps,"
            " 5) a 30-60-90 day learning plan. Keep it concise."
        )
        user = (
            "CV JSON:\n" + cv_preview + "\n\n"
            "Relevant Career Tree Context (do not hallucinate outside of this):\n\n"
            + context
            + "\n\nReturn a short structured list with bullets."
        )
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ]
        completion = chat.complete_chat(messages)
        return completion.choices[0].message.content or ""
    finally:
        if own_model is not None:
            own_model.unload()


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract a CV and recommend roles using the Career Tree RAG")
    parser.add_argument("--cv", required=True, help="Path to the CV file (.pdf/.docx/.txt/.md)")
    parser.add_argument("--model", default=None, help="Foundry Local model alias for recommendations")
    parser.add_argument(
        "--extract-model",
        default=None,
        help="Foundry Local model alias for CV extraction (defaults to --model or built-in default)",
    )
    parser.add_argument("--top-k", type=int, default=5, help="How many roles to retrieve")
    parser.add_argument("--layer", default=None, help="Restrict retrieval to a specific layer")
    parser.add_argument("--extract-only", action="store_true", help="Only extract JSON and print it")
    parser.add_argument("--no-cache", action="store_true", help="Ignore cached CV extraction results")
    args = parser.parse_args()

    cv_path = Path(args.cv)
    if not cv_path.exists():
        raise SystemExit(f"CV not found: {cv_path}")

    print("Starting career-rag-recommend…")
    print("Reading CV…")
    extract_model = args.extract_model or args.model or DEFAULT_CV_EXTRACTION_MODEL
    answer_model = args.model or extract_model
    use_cache = not args.no_cache

    if args.extract_only:
        print("Extracting structured data (loading local LLM; first run can take a minute)…")
        cv_json = load_or_extract_cv(cv_path, model_alias=extract_model, use_cache=use_cache)
        print(json.dumps(cv_json, ensure_ascii=False, indent=2))
        return

    from rag.llm import load_model, make_chat

    print("Loading local LLM (first run can take a minute)…")
    model = load_model(answer_model)
    try:
        print("Extracting structured data…")
        extract_chat = make_chat(model, temperature=0.1, max_tokens=2000)
        cv_json = load_or_extract_cv(
            cv_path,
            model_alias=extract_model,
            chat=extract_chat,
            use_cache=use_cache,
        )
        print(json.dumps(cv_json, ensure_ascii=False, indent=2))

        print("Loading vector search (first run may take 20-40s)…")
        from rag.retrieve import format_context

        print("Retrieving matching roles…")
        result = recommend_from_cv(cv_json, top_k=args.top_k, layer=args.layer)
        hits = result["hits"]
        print("=== Retrieved Context ===\n")
        print(format_context(hits) or "(no sufficiently relevant context)")

        if not hits:
            print("\n(no recommendations — not enough relevant context)")
            return

        print("\n=== Recommendations ===\n")
        answer_chat = make_chat(model, temperature=0.2, max_tokens=1200)
        plan = _plan_recommendations(cv_json, hits, chat=answer_chat)
        print(plan)
    finally:
        model.unload()


if __name__ == "__main__":
    main()
